"""Part 1: Document setup and helper functions"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_code_block(doc, code, lang="typescript"):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)

def add_info_box(doc, text, label="NOTE"):
    p = doc.add_paragraph()
    r = p.add_run(f"[{label}] ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def build_title_page(doc):
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("TRUVA PROTOCOL")
    r.font.size = Pt(36)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("Solana Frontier Hackathon — Complete Submission & Implementation Guide")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
    
    doc.add_paragraph()
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = d.add_run("All 13 Applicable Tracks · Code Changes · Submission Strategy")
    r3.font.size = Pt(12)
    r3.font.italic = True
    
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = info.add_run("Deadline: ~May 12, 2026 | Total Prize Pool: $437,914 | 54 Tracks")
    r4.font.size = Pt(11)
    r4.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    r4.bold = True
    doc.add_page_break()
